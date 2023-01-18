import re

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


class Passport:
    def __init__(self):
        self.passport_number_title = "№ реставраційного паспорта:\n\n"
        self.inventory_number_title = "інвентарний № пам'ятки\n\n"
        self.acceptance_number_title = "Акт приймання\n"
        self.ministry_title = "\nМіністерство культури та інформаційної політики України"
        self.passport_type_title = "ПАСПОРТ РЕСТАВРАЦІЇ ПАМ'ЯТКИ ІСТОРІЇ ТА КУЛЬТУРИ (РУХОМОЇ)"
        self.institution_title = "назва закладу, який здійснює реставрацію"
        self.department_title = "назва відділу/сектору"
        self.typological_title = "1. Типологічна приналежність пам'ятки"
        self.definition_title = "Визначення характер пам'ятки:"
        self.fine_art_title = "пам'ятка образотворчого мистецтва:"
        self.applied_art_title = "пам'ятка декоративно-ужиткового мистецтва:"
        self.archeological_title = "археологічна пам'ятка:"
        self.documentary_title = "документальна пам'ятка:"
        self.other_object_title = "інша пам'ятка історії та культури:"
        self.object_owner_title = "2. Місце постійного зберігання, власник пам'ятки"
        self.attribution_data_from_acceptance_report_title = "3. Атрибутивні дані про пам'ятку згідно з актом приймання"
        self.clarification_title = "Уточнення в процесі реставрації"
        self.author_title = "Автор: "
        self.object_name_title = "Назва: "
        self.time_of_creation_title = "Час створення: "
        self.material_title = "Матеріал, основа: "
        self.technique_title = "Техніка виконання: "
        self.object_size_title = "Розміри: "
        self.weight_title = "Вага: "
        self.reason_title = "4. Підстава для проведення реставраційних заходів"
        self.dates_title = "Дати"
        self.restorers_title = "Реставратори"
        self.object_input_date_title = "Дата передання: "
        self.execute_restorer_title = "Виконавець: "
        self.object_output_date_title = "Дата завершення: "
        self.responsible_restorer_title = "Керівник: "
        self.origin_description_title = '''5. Основні дані з історії пам'ятки (довідка про побутування; 
відомості про умови зберігання, попередні дослідження, консерваційно-реставраційні заходи тощо), 
джерело надходження інформації'''
        self.condition_before_restoration_title = "6. Стан пам'ятки до реставрації"
        self.appearance_description_title = "6.1 За візуальним спостереженням:"
        self.damages_description_title = "6.1.2 Втрати та пошкодження:"
        self.signs_description_title = "6.1.3 Старі номери та позначення:"
        self.size_description_title = "6.1.4 Розміри:"
        self.researches_title = "6.2. За даними лабораторних досліджень:"
        self.purposes_researches_title = "Мета дослідження"
        self.methods_researches_title = "Методи і результати дослідження"
        self.executor_date_researches_title = "Виконавець та дата"
        self.results_researches_title = "6.3 Загальний висновок за результатами досліджень:"
        self.restoration_program_title = "7. Програма проведення реставраційних заходів та їх обгрунтування:"
        self.restoration_program_title = "7. Програма проведення реставраційних заходів та їх обгрунтування:"
        self.sequence_of_treatments_title = "Послідовність заходів"
        self.program_approved_meeting_title = '''Програма затверджена на засіданні науково-реставраційної/реставраційної ради
____________________________________________________________________________________________________
місце проведення ради (назва закладу)'''
        self.approved_meeting_protocol_number_title = "Протокол №____ від '     '_______________ 20___р."
        self.appointed_responsible_restorer_title = "Керівником роботи призначено (ПІБ):___________________________________________________________"
        self.meeting_secretary_title = "Голова або секретар науково-реставраційної/реставраційної ради (ПІБ):__________________________________________________________________________________________________"
        self.reapproved_restoration_program_title = "8. Зміни в програмі реставраційних заходів та їх обгрунтування:\n"
        self.reapproved_program_meeting_place_title = '''Зміни в програмі затверджені на засіданні науково-реставраційної/реставраційної ради
_______________________________________________________________________________________________________
місце проведення ради (назва закладу)'''
        self.reapproved_meeting_protocol_number_title = "Протокол №____ від '     '_______________ 20___р."
        self.treatments_title = "9. Проведення реставраційних заходів:"
        self.treatments_descriptions_title = "Опис операцій із зазначенням методу, методики, технології, інструментарію"
        self.treatments_chemicals_title = "Матеріали, хімікати (концентрація %)"
        self.treatments_executor_date_title = "Виконавець та дата"
        self.treatments_results_title = "10. Стислий опис реставраційних заходів; опис змін технічного та " \
                                        "зовнішнього стану пам'ятки після реставрації, уточнення атрибуції тощо:"
        self.responsible_restorer_signature_title = "Керівник роботи (ПІБ та підпис):"
        self.signature_place = '_______________________________________'
        self.execute_restorer_signature_title = "Виконавець (ПІБ та підпис):"
        self.meeting_conclusion_title = "\n\n\n\n\n\n11. Висновок науково-реставраційної/реставраційної ради (витяг з протоколу)"
        self.meeting_place_title = '''
____________________________________________________________________________________________________
____________________________________________________________________________________________________
____________________________________________________________________________________________________
місце проведення ради (назва закладу)'''
        self.after_restoration_meeting_protocol_number_title = "Протокол №____ від '     '_______________ 20___р."
        self.recommendations_title = "12. Рекомендації щодо умов зберігання пам'ятки"
        self.illustrative_material_title = "13. Ілюстративний матеріал:"
        self.before_restoration_image_description_title = 'Мал.1. "До реставрації": '
        self.process_restoration_image_description_title = 'Мал.2. "В процесі реставрації": '
        self.after_restoration_image_description_title = 'Мал.3. "Після реставрації"'
        self.appendix_title = '14. Додатки до паспорта: '
        self.related_materials_title = 'ілюстративний матеріал, результати дослідження тощо'
        self.photos_title = "фотовідбитки"
        self.photos_amount = ''
        self.units_title = "од."
        self.map_schemes_title = "картосхеми"
        self.map_schemes_amount = ''
        self.additional_researches_results_title = 'результати дослідження'
        self.additional_researches_amount = ''
        self.other_related_materials_title = 'інше (зазначити)'
        self.other_related_materials_amount = ''
        self.attachment_illustrative_material_place_title = '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n' \
                                                            'місце клапана для\nматеріалів додатка\n\n\n'
        self.object_transferred_place_title = '''Після реставрації пам'ятка передана в:
____________________________________________________________________________________________________
назва організації, № і дата акту про передачу'''
        self.responsible_treasurer_and_storage_place_title = '''
____________________________________________________________________________________________________
назва фондової групи, відповідальний хранитель'''
        self.give_back_report_title = "Акт повернення №____ від '     '_______________ 20___р."
        self.passport_copies_transferred_place_title = '''
Копії паспорту в 2-х прим. передані в:
____________________________________________________________________________________________________
назва організації, № накладної і дата передачі паспортів'''
        self.head_of_institution_title = 'Керівник організації:'
        self.head_of_the_department_title = 'Завідувач відділу:'
        self.head_of_restoration_title = 'Керівник роботи:'
        self.executor_of_restoration_title = 'Виконавець роботи:'
        self.restorers_and_other_executors_title = 'Реставратори та інші виконавці:'
        self.empty_row_place = "___________________________________________________________________________"
        self.other_executors_signature_title = '''
М.П._______________________________________________________________________________________________
____________________________________________________________________________________________________
____________________________________________________________________________________________________
(ПІБ), посада, кваліфікаційна категорія, підпис'''
        self.condition_check_date_title = 'Дата огляду'
        self.condition_description_title = "Стан пам'ятки"
        self.storage_condition_title = 'Умови зберігання'
        self.checking_person_signature_title = 'Посада, ПІБ, підпис'
        self.indent = '\n\n\n\n\n\n\n\n\n\n'

    def build_passport(self, passport_number, inventory_number, acceptance_number, institution_name, department_name,
                       definition, typological, object_owner, author, clarified_author, object_name,
                       clarified_object_name, time_of_creation, clarified_time_of_creation, material,
                       clarified_material, technique, clarified_technique, object_size, clarified_size, weight,
                       clarified_weight, reason,
                       object_input_date, execute_restorer, object_output_date, responsible_restorer,
                       origin_description, appearance_description, damages_description, signs_description,
                       size_description, purposes_researches, methods_researches, executor_date_researches,
                       results_researches, restoration_program, treatments_descriptions, treatments_chemicals,
                       treatments_executor_date, treatments_results):

        document = Document()

        passport_identity_table = document.add_table(rows=1, cols=3)
        passport_identity_table.autofit = False
        passport_identity_table.allow_autofit = False
        passport_identity_table.columns[0].width = Inches(1.6)
        passport_identity_table.columns[1].width = Inches(1.4)
        passport_identity_table.columns[2].width = Inches(1.3)
        passport_identity_table.style = 'TableGrid'
        passport_identity_table.alignment = 2

        passport_identity_parameter = passport_identity_table.rows[0].cells
        passport_identity_parameter[0].text = f"{self.passport_number_title}{passport_number}"
        passport_identity_parameter[1].text = f"{self.inventory_number_title}{inventory_number}"
        passport_identity_parameter[2].text = f"{self.acceptance_number_title}{acceptance_number}"

        font_styles = document.styles
        font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        font_object = font_charstyle.font
        font_object.name = 'Times New Roman'
        font_object.size = Pt(10)

        ministry_paragraph = document.add_paragraph()
        ministry_paragraph.add_run(f"{self.ministry_title}", style='CommentsStyle').bold = True
        ministry_paragraph.alignment = 1

        passport_type_paragraph = document.add_paragraph()
        passport_type_paragraph.add_run(f"{self.passport_type_title}", style='CommentsStyle').bold = True
        passport_type_paragraph.alignment = 1

        institution_paragraph = document.add_paragraph()
        institution_paragraph.add_run(f'{institution_name}', style='CommentsStyle').bold = True
        institution_paragraph.alignment = 1

        institution_parameter = document.add_paragraph()
        institution_parameter.add_run(f"{self.institution_title}")
        institution_parameter.alignment = 1

        department_paragraph = document.add_paragraph()
        department_paragraph.add_run(f'{department_name}', style='CommentsStyle').bold = True
        department_paragraph.alignment = 1

        department_parameter = document.add_paragraph()
        department_parameter.add_run(f'{self.department_title}')
        department_parameter.alignment = 1

        typological_paragraph = document.add_paragraph()
        typological_paragraph.add_run(f'{self.typological_title}').bold = True
        typological_paragraph.alignment = 0

        typological_table = document.add_table(rows=2, cols=6)
        typological_table.style = 'TableGrid'
        typological_table.alignment = 1

        typological_parameter = typological_table.rows[0].cells
        typological_parameter[0].text = f'{self.definition_title}'
        typological_parameter[1].text = f'{self.fine_art_title}'
        typological_parameter[2].text = f'{self.applied_art_title}'
        typological_parameter[3].text = f'{self.archeological_title}'
        typological_parameter[4].text = f'{self.documentary_title}'
        typological_parameter[5].text = f'{self.other_object_title}'

        typological_input = typological_table.rows[1].cells
        typological_input[0].text = f'{definition}'
        if typological == '(1)':
            typological_input[1].text = f'{typological}'
        elif typological == '(2)':
            typological_input[2].text = f'{typological}'
        elif typological == '(3)':
            typological_input[3].text = f'{typological}'
        elif typological == '(4)':
            typological_input[4].text = f'{typological}'
        elif typological == '(5)':
            typological_input[5].text = f'{typological}'
        else:
            typological_input[5].text = ''

        object_owner_paragraph = document.add_paragraph()
        object_owner_paragraph.add_run(f'{self.object_owner_title}').bold = True
        object_owner_paragraph.alignment = 1

        object_owner_input = document.add_paragraph()
        object_owner_input.add_run(f'{object_owner}')
        object_owner_input.alignment = 1

        attribution_data_table = document.add_table(rows=8, cols=2)
        attribution_data_table.style = 'TableGrid'
        attribution_data_table.alignment = 1

        attribution_data_parameter = attribution_data_table.rows[0].cells
        attribution_data_parameter[0].text = f'{self.attribution_data_from_acceptance_report_title}'
        attribution_data_parameter[1].text = f'{self.clarification_title}'

        author_parameter = attribution_data_table.rows[1].cells
        author_parameter[0].text = f'{self.author_title}{author}'
        author_parameter[1].text = f'{clarified_author}'

        object_name_parameter = attribution_data_table.rows[2].cells
        object_name_parameter[0].text = f'{self.object_name_title}{object_name}'
        object_name_parameter[1].text = f'{clarified_object_name}'

        time_of_creation_parameter = attribution_data_table.rows[3].cells
        time_of_creation_parameter[0].text = f'{self.time_of_creation_title}{time_of_creation}'
        time_of_creation_parameter[1].text = f'{clarified_time_of_creation}'

        material_parameter = attribution_data_table.rows[4].cells
        material_parameter[0].text = f'{self.material_title}{material}'
        material_parameter[1].text = f'{clarified_material}'

        technique_parameter = attribution_data_table.rows[5].cells
        technique_parameter[0].text = f'{self.technique_title}{technique}'
        technique_parameter[1].text = f'{clarified_technique}'

        object_size_parameter = attribution_data_table.rows[6].cells
        object_size_parameter[0].text = f'{self.object_size_title}{object_size}'
        object_size_parameter[1].text = f'{clarified_size}'

        weight_parameter = attribution_data_table.rows[7].cells
        weight_parameter[0].text = f'{self.weight_title }{weight}'
        weight_parameter[1].text = f'{clarified_weight}'

        reason_paragraph = document.add_paragraph()
        reason_paragraph.add_run(f'{self.reason_title}').bold = True
        reason_paragraph.alignment = 1

        reason_input = document.add_paragraph()
        reason = re.sub(r'<br>', '\n', reason)
        reason_input.add_run(f'{reason}')
        reason_input.paragraph_format.first_line_indent = Pt(18)
        reason_input.alignment = 0

        dates_and_restorers_table = document.add_table(rows=3, cols=2)
        dates_and_restorers_table.style = 'TableGrid'
        dates_and_restorers_table.alignment = 1

        dates_and_restorers_parameter = dates_and_restorers_table.rows[0].cells
        dates_and_restorers_parameter[0].text = f'{self.dates_title}'
        dates_and_restorers_parameter[1].text = f'{self.restorers_title}'

        dates_and_restorers_parameter = dates_and_restorers_table.rows[1].cells
        dates_and_restorers_parameter[0].text = f'{self.object_input_date_title}{object_input_date}'
        dates_and_restorers_parameter[1].text = f'{self.execute_restorer_title}{execute_restorer}'

        dates_and_restorers_parameter = dates_and_restorers_table.rows[2].cells
        dates_and_restorers_parameter[0].text = f'{self.object_output_date_title}{object_output_date}'
        dates_and_restorers_parameter[1].text = f'{self.responsible_restorer_title}{responsible_restorer}'

        document.add_page_break()

        origin_description_parameter = document.add_paragraph()
        origin_description_parameter.add_run(f'{self.origin_description_title}').bold = True
        origin_description_parameter.alignment = 0

        origin_description_input = document.add_paragraph()
        origin_description = re.sub(r'<br>', '\n', origin_description)
        origin_description_input.add_run(f'{origin_description}')
        origin_description_input.paragraph_format.first_line_indent = Pt(18)
        origin_description_input.alignment = 0

        appearance_description_paragraph = document.add_paragraph()
        appearance_description_paragraph.add_run(f'{self.condition_before_restoration_title}').bold = True
        appearance_description_paragraph.alignment = 0
        visual_observation_title = document.add_paragraph()
        visual_observation_title.add_run(f'{self.appearance_description_title}', 'Emphasis')
        visual_observation_title.alignment = 0
        appearance_description_input = document.add_paragraph()
        appearance_description = re.sub(r'<br>', '\n', appearance_description)
        appearance_description_input.add_run(f'{appearance_description}')
        appearance_description_input.paragraph_format.first_line_indent = Pt(18)
        appearance_description_input.alignment = 0

        damages_description_paragraph = document.add_paragraph()
        damages_description_paragraph.add_run(f'{self.damages_description_title}', 'Emphasis')
        damages_description_paragraph.alignment = 0

        damages_description_input = document.add_paragraph()
        damages_description = re.sub(r'<br>', '\n', damages_description)
        damages_description_input.add_run(f'{damages_description}')
        damages_description_input.paragraph_format.first_line_indent = Pt(18)
        damages_description_input.alignment = 0

        signs_description_paragraph = document.add_paragraph()
        signs_description_paragraph.add_run(f'{self.signs_description_title}', 'Emphasis')
        signs_description_paragraph.alignment = 0

        signs_description_input = document.add_paragraph()
        signs_description = re.sub(r'<br>', '\n', signs_description)
        signs_description_input.add_run(f'{signs_description}')
        signs_description_input.paragraph_format.first_line_indent = Pt(18)
        signs_description_input.alignment = 0

        size_description_paragraph = document.add_paragraph()
        size_description_paragraph.add_run(f'{self.size_description_title}', 'Emphasis')
        size_description_paragraph.alignment = 0

        size_description_input = document.add_paragraph()
        size_description = re.sub(r'<br>', '\n', size_description)
        size_description_input.add_run(f'{size_description}')
        size_description_input.paragraph_format.left_indent = Inches(0.25)
        size_description_input.alignment = 0

        document.add_page_break()

        researches_paragraph = document.add_paragraph()
        researches_paragraph.add_run(f'{self.researches_title}').bold = True
        researches_paragraph.alignment = 0

        researches_table = document.add_table(rows=2, cols=3)
        researches_table.autofit = False
        researches_table.allow_autofit = False
        researches_table.columns[0].width = Inches(2.0)
        researches_table.columns[1].width = Inches(3.0)
        researches_table.columns[2].width = Inches(1.0)
        researches_table.style = 'TableGrid'
        researches_table.alignment = 1

        researches_parameter = researches_table.rows[0].cells
        researches_parameter[0].text = f'{self.purposes_researches_title}'
        researches_parameter[1].text = f'{self.methods_researches_title}'
        researches_parameter[2].text = f'{self.executor_date_researches_title}'

        researches_parameter = researches_table.rows[1].cells
        purposes_researches = re.sub(r'<br>', '\n', purposes_researches)
        researches_parameter[0].text = f'{purposes_researches}'
        methods_researches = re.sub(r'<br>', '\n', methods_researches)
        researches_parameter[1].text = f'{methods_researches}'
        executor_date_researches = re.sub(r'<br>', '\n', executor_date_researches)
        researches_parameter[2].text = f'{executor_date_researches}'

        results_researches_paragraph = document.add_paragraph()
        results_researches_paragraph.add_run(f'{self.results_researches_title}').bold = True
        results_researches_paragraph.alignment = 0

        results_researches_input = document.add_paragraph()
        results_researches = re.sub(r'<br>', '\n', results_researches)
        results_researches_input.add_run(f'{results_researches}')
        results_researches_input.paragraph_format.first_line_indent = Pt(18)
        results_researches_input.alignment = 0

        document.add_page_break()

        restoration_program_paragraph = document.add_paragraph()
        restoration_program_paragraph.add_run(f'{self.restoration_program_title}').bold = True
        restoration_program_paragraph.alignment = 0
        treatments_order_program_paragraph = document.add_paragraph()
        treatments_order_program_paragraph.add_run(f'{self.sequence_of_treatments_title}', 'Emphasis')
        treatments_order_program_paragraph.alignment = 0
        restoration_program_input = document.add_paragraph()
        restoration_program = re.sub(r'<br>', '\n', restoration_program)
        restoration_program_input.add_run(f'{restoration_program}')
        restoration_program_input.alignment = 0

        program_approved_meeting_table = document.add_table(rows=2, cols=1)
        program_approved_meeting_table.style = 'TableGrid'
        program_approved_meeting_table.alignment = 1

        program_approved_meeting_parameter = program_approved_meeting_table.rows[0].cells
        program_approved_meeting_parameter[0].text = f'{self.program_approved_meeting_title}'

        program_approved_meeting_parameter = program_approved_meeting_table.rows[1].cells
        program_approved_meeting_parameter[0].text = f'{self.approved_meeting_protocol_number_title}'

        appointed_responsible_restorer_paragraph = document.add_paragraph()
        appointed_responsible_restorer_paragraph.add_run(f'{self.appointed_responsible_restorer_title}')
        appointed_responsible_restorer_paragraph.alignment = 0

        meeting_secretary_paragraph = document.add_paragraph()
        meeting_secretary_paragraph.add_run(f'{self.meeting_secretary_title}')
        meeting_secretary_paragraph.alignment = 0

        reapproved_restoration_program_paragraph = document.add_paragraph()
        reapproved_restoration_program_paragraph.add_run(f'{self.reapproved_restoration_program_title}').bold = True
        reapproved_restoration_program_paragraph.alignment = 0

        reapproved_program_meeting_table = document.add_table(rows=2, cols=1)
        reapproved_program_meeting_table.style = 'TableGrid'
        reapproved_program_meeting_table.alignment = 1

        reapproved_program_meeting_parameter = reapproved_program_meeting_table.rows[0].cells
        reapproved_program_meeting_parameter[0].text = f'{self.reapproved_program_meeting_place_title}'

        reapproved_program_meeting_parameter = reapproved_program_meeting_table.rows[1].cells
        reapproved_program_meeting_parameter[0].text = f'{self.reapproved_meeting_protocol_number_title}'

        reapproved_meeting_secretary_paragraph = document.add_paragraph()
        reapproved_meeting_secretary_paragraph.add_run(f'{self.meeting_secretary_title}')
        reapproved_meeting_secretary_paragraph.alignment = 0

        document.add_page_break()

        treatments_paragraph = document.add_paragraph()
        treatments_paragraph.add_run(f'{self.treatments_title}').bold = True
        treatments_paragraph.alignment = 0

        treatments_table = document.add_table(rows=2, cols=3)
        treatments_table.autofit = False
        treatments_table.allow_autofit = False
        treatments_table.columns[0].width = Inches(3.0)
        treatments_table.columns[1].width = Inches(1.6)
        treatments_table.columns[2].width = Inches(1.0)
        treatments_table.style = 'TableGrid'
        treatments_table.alignment = 1

        treatments_parameter = treatments_table.rows[0].cells
        treatments_parameter[0].text = f'{self.treatments_descriptions_title}'
        treatments_parameter[1].text = f'{self.treatments_chemicals_title}'
        treatments_parameter[2].text = f'{self.treatments_executor_date_title}'

        treatments_parameter = treatments_table.rows[1].cells
        treatments_descriptions = re.sub(r'<br>', '\n', treatments_descriptions)
        treatments_parameter[0].text = f'{treatments_descriptions}'
        treatments_chemicals = re.sub(r'<br>', '\n', treatments_chemicals)
        treatments_parameter[1].text = f'{treatments_chemicals}'
        treatments_executor_date = re.sub(r'<br>', '\n', treatments_executor_date)
        treatments_parameter[2].text = f'{treatments_executor_date}'

        document.add_page_break()

        treatments_paragraph = document.add_paragraph()
        treatments_paragraph.add_run(f'{self.treatments_results_title}').bold = True
        treatments_paragraph.alignment = 0

        treatments_results_input = document.add_paragraph()
        treatments_results = re.sub(r'<br>', '\n', treatments_results)
        treatments_results_input.add_run(f'{treatments_results}')
        treatments_results_input.paragraph_format.first_line_indent = Pt(18)
        treatments_results_input.alignment = 0

        result_responsible_restorer_paragraph = document.add_paragraph()
        result_responsible_restorer_paragraph.add_run(f'{self.responsible_restorer_signature_title}{responsible_restorer}{self.signature_place}')
        result_responsible_restorer_paragraph.alignment = 2

        result_execute_restorer_paragraph = document.add_paragraph()
        result_execute_restorer_paragraph.add_run(f'{self.execute_restorer_signature_title}{execute_restorer}{self.signature_place}')
        result_execute_restorer_paragraph.alignment = 2

        result_object_output_date_paragraph = document.add_paragraph()
        result_object_output_date_paragraph.add_run(f'{object_output_date}')
        result_object_output_date_paragraph.alignment = 2

        meeting_conclusion_paragraph = document.add_paragraph()
        meeting_conclusion_paragraph.add_run(f'{self.meeting_conclusion_title}').bold = True
        meeting_conclusion_paragraph.alignment = 0

        meeting_conclusion_table = document.add_table(rows=2, cols=1)
        meeting_conclusion_table.style = 'TableGrid'
        meeting_conclusion_table.alignment = 1

        meeting_conclusion_parameter = meeting_conclusion_table.rows[0].cells
        meeting_conclusion_parameter[0].text = f'{self.meeting_place_title}'

        meeting_conclusion_parameter = meeting_conclusion_table.rows[1].cells
        meeting_conclusion_parameter[0].text = f'{self.after_restoration_meeting_protocol_number_title}'

        conclusion_meeting_secretary_paragraph = document.add_paragraph()
        conclusion_meeting_secretary_paragraph.add_run(f'{self.meeting_secretary_title}')
        conclusion_meeting_secretary_paragraph.alignment = 0

        recommendations_paragraph = document.add_paragraph()
        recommendations_paragraph.add_run(self.recommendations_title).bold = True
        recommendations_paragraph.alignment = 0

        recommendations_table = document.add_table(rows=5, cols=1)
        recommendations_table.style = 'TableGrid'
        recommendations_table.alignment = 1

        recommendations_execute_restorer_paragraph = document.add_paragraph()
        recommendations_execute_restorer_paragraph.add_run(f'{self.execute_restorer_signature_title}{execute_restorer}{self.signature_place}')
        recommendations_execute_restorer_paragraph.alignment = 2

        recommendations_output_date_paragraph = document.add_paragraph()
        recommendations_output_date_paragraph.add_run(f'{object_output_date}')
        recommendations_output_date_paragraph.alignment = 2

        document.add_page_break()

        illustrative_material_paragraph = document.add_paragraph()
        illustrative_material_paragraph.add_run(f'{self.illustrative_material_title}').bold = True
        illustrative_material_paragraph.alignment = 0

        document.add_page_break()

        counted_illustrative_material_paragraph = document.add_paragraph()
        counted_illustrative_material_paragraph.add_run(f'{self.appendix_title}').bold = True
        counted_illustrative_material_paragraph.add_run(f'{self.related_materials_title}')
        counted_illustrative_material_paragraph.alignment = 0

        counted_illustrative_material_table = document.add_table(rows=4, cols=3)
        counted_illustrative_material_table.autofit = False
        counted_illustrative_material_table.allow_autofit = False
        counted_illustrative_material_table.columns[0].width = Inches(1.5)
        counted_illustrative_material_table.columns[1].width = Inches(1.0)
        counted_illustrative_material_table.columns[2].width = Inches(0.5)
        counted_illustrative_material_table.style = 'TableGrid'
        counted_illustrative_material_table.alignment = 0

        counted_illustrative_material_parameter = counted_illustrative_material_table.rows[0].cells
        counted_illustrative_material_parameter[0].text = f'{self.photos_title}'
        counted_illustrative_material_parameter[1].text = f'{self.photos_amount}'
        counted_illustrative_material_parameter[2].text = f'{self.units_title}'

        counted_illustrative_material_parameter = counted_illustrative_material_table.rows[1].cells
        counted_illustrative_material_parameter[0].text = f'{self.map_schemes_title}'
        counted_illustrative_material_parameter[1].text = f'{self.map_schemes_amount}'
        counted_illustrative_material_parameter[2].text = f'{self.units_title}'

        counted_illustrative_material_parameter = counted_illustrative_material_table.rows[2].cells
        counted_illustrative_material_parameter[0].text = f'{self.additional_researches_results_title}'
        counted_illustrative_material_parameter[1].text = f'{self.additional_researches_amount}'
        counted_illustrative_material_parameter[2].text = f'{self.units_title}'

        counted_illustrative_material_parameter = counted_illustrative_material_table.rows[3].cells
        counted_illustrative_material_parameter[0].text = f'{self.other_related_materials_title}'
        counted_illustrative_material_parameter[1].text = f'{self.other_related_materials_amount}'
        counted_illustrative_material_parameter[2].text = f'{self.units_title}'

        attachment_illustrative_material_place_paragraph = document.add_paragraph()
        attachment_illustrative_material_place_paragraph.add_run(f'{self.attachment_illustrative_material_place_title}')
        attachment_illustrative_material_place_paragraph.alignment = 2

        document.add_page_break()

        object_transferred_table = document.add_table(rows=4, cols=1)
        object_transferred_table.style = 'TableGrid'
        object_transferred_table.alignment = 1

        object_transferred_parameter = object_transferred_table.rows[0].cells
        object_transferred_parameter[0].text = f'{self.object_transferred_place_title}'

        object_transferred_parameter = object_transferred_table.rows[1].cells
        object_transferred_parameter[0].text = f'{self.responsible_treasurer_and_storage_place_title}'

        object_transferred_parameter = object_transferred_table.rows[2].cells
        object_transferred_parameter[0].text = f'{self.give_back_report_title}'

        object_transferred_parameter = object_transferred_table.rows[3].cells
        object_transferred_parameter[0].text = f'{self.passport_copies_transferred_place_title}'

        first_divider_paragraph = document.add_paragraph()
        first_divider_paragraph.add_run('\n')
        first_divider_paragraph.alignment = 1

        related_employees_table = document.add_table(rows=5, cols=2)
        related_employees_table.autofit = False
        related_employees_table.allow_autofit = False
        related_employees_table.columns[0].width = Inches(1.7)
        related_employees_table.columns[1].width = Inches(4.4)
        related_employees_table.style = 'TableGrid'
        related_employees_table.alignment = 1

        related_employees_parameter = related_employees_table.rows[0].cells
        related_employees_parameter[0].text = f'{self.head_of_institution_title}'

        related_employees_parameter = related_employees_table.rows[1].cells
        related_employees_parameter[0].text = f'{self.head_of_the_department_title}'

        related_employees_parameter = related_employees_table.rows[2].cells
        related_employees_parameter[0].text = f'{self.head_of_restoration_title}'
        related_employees_parameter[1].text = f'{responsible_restorer}'

        related_employees_parameter = related_employees_table.rows[3].cells
        related_employees_parameter[0].text = f'{self.executor_of_restoration_title}'
        related_employees_parameter[1].text = f'{execute_restorer}'

        related_employees_parameter = related_employees_table.rows[4].cells
        related_employees_parameter[0].text = f'{self.restorers_and_other_executors_title}'
        related_employees_parameter[1].text = f'{self.empty_row_place}'

        second_divider_paragraph = document.add_paragraph()
        second_divider_paragraph.add_run('\n')
        second_divider_paragraph.alignment = 1

        other_executors_table = document.add_table(rows=1, cols=1)
        other_executors_table.style = 'TableGrid'
        other_executors_table.alignment = 1

        other_executors_parameter = other_executors_table.rows[0].cells
        other_executors_parameter[0].text = f'{self.other_executors_signature_title}'

        third_divider_paragraph = document.add_paragraph()
        third_divider_paragraph.add_run('\n')
        third_divider_paragraph.alignment = 1

        condition_observation_table = document.add_table(rows=2, cols=4)
        condition_observation_table.style = 'TableGrid'
        condition_observation_table.alignment = 1

        condition_observation_parameter = condition_observation_table.rows[0].cells
        condition_observation_parameter[0].text = f'{self.condition_check_date_title}'
        condition_observation_parameter[1].text = f'{self.condition_description_title}'
        condition_observation_parameter[2].text = f'{self.storage_condition_title}'
        condition_observation_parameter[3].text = f'{self.checking_person_signature_title}'

        condition_observation_parameter = condition_observation_table.rows[1].cells
        condition_observation_parameter[0].text = f'{self.indent}'

        document.add_page_break()

        document.save('filled_passport.docx')

# For Linux Ubuntu Apache2 replace last line by this:
#       document.save('/var/www/html/yourProjectDirectory/filled_passport.docx')
