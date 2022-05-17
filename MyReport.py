# 10
import sqlite3
import os
from flask import Flask, render_template, url_for, request, flash, session, redirect, abort, g, send_file, make_response
from flask import send_from_directory
from werkzeug.utils import secure_filename
from FDataBase import FDataBase
from heritageLogic import GeneralHeritageDescription, Materials
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import re
import jyserver.Flask as jsf

DATABASE = '/tmp/flsite.db'
DEBUG = True
SECRET_KEY = 'fdgdfgdfggf786hfg6hfg6h7f'
ALLOWED_EXTENSIONS = {'png'}
UPLOAD_FOLDER = 'C:/Users/ASUS/Desktop/pyCharm_projects/restoreConrol/static/Intelligible_illustrations'
MAX_CONTENT_LENGTH = 16 * 1000 * 1000
app = Flask(__name__)
app.config.from_object(__name__)
app.config.update(dict(DATABASE=os.path.join(app.root_path, 'flsite,db')))


@jsf.use(app)
class ChangeFormsForMaterial:
    def __init__(self, initial_changed_value='', fillInExamples=GeneralHeritageDescription.AppearanceDescription(),
                 fillInCeramics=Materials.Ceramics(), fillInIron=Materials.Iron(), fillInCuprum=Materials.Cuprum(),
                 fillInSilver=Materials.Silver(), fillInWood=Materials.Wood(), fillInMetal=Materials.Metal()):
        # General texts for any heritage object
        self.changed_appearance_description = fillInExamples.general_appearance_description
        self.changed_damages_description = fillInExamples.general_damages_description
        self.changed_ConductTitle, self.changed_ConductDescription, self.changed_restoration_program, \
        self.changed_treatments_descriptions, self.changed_treatments_chemicals = initial_changed_value, \
                                                                                  initial_changed_value, initial_changed_value, initial_changed_value, initial_changed_value
        # Ceramics (wide material field)
        self.ceramics_appearance_description = fillInCeramics.ceramics_appearance_description
        self.ceramics_damage_description = fillInCeramics.ceramics_damage_description
        # Wood (wide material field)
        self.wood_appearance_description = fillInWood.wood_appearance_description
        self.wood_damage_description = fillInWood.wood_damage_description
        # Metal (wide material field)
        self.metal_appearance_description = fillInMetal.metal_appearance_description
        self.metal_damage_description = fillInMetal.metal_damage_description
        # Iron
        self.ironConductTitle = fillInIron.iron_research_titles
        self.ironConductDescription = fillInIron.iron_research_descriptions
        self.restoration_program_iron = fillInIron.iron_restoration_program
        self.iron_treatments_descriptions = fillInIron.iron_treatments_descriptions
        self.iron_treatments_chemicals = fillInIron.iron_treatments_chemicals
        # Cuprum
        self.cuprumConductTitle = fillInCuprum.cuprum_research_titles
        self.cuprumConductDescription = fillInCuprum.cuprum_research_descriptions
        self.restoration_program_cuprum = fillInCuprum.cuprum_restoration_program
        self.cuprum_treatments_descriptions = fillInCuprum.cuprum_treatments_descriptions
        self.cuprum_treatments_chemicals = fillInCuprum.cuprum_treatments_chemicals
        # Silver
        self.silverConductTitle = fillInSilver.silver_research_titles
        self.silverConductDescription = fillInSilver.silver_research_descriptions
        self.restoration_program_silver = fillInSilver.silver_restoration_program
        self.silver_treatments_descriptions = fillInSilver.silver_treatments_descriptions
        self.silver_treatments_chemicals = fillInSilver.silver_treatments_chemicals

    def addCeramics(self):
        self.changed_appearance_description += self.ceramics_appearance_description
        self.js.document.getElementById(
            "changed_appearance_description").innerHTML = self.changed_appearance_description
        self.changed_damages_description += self.ceramics_damage_description
        self.js.document.getElementById("changed_damages_description").innerHTML = self.changed_damages_description

    def addWood(self):
        self.changed_appearance_description += self.wood_appearance_description
        self.js.document.getElementById(
            "changed_appearance_description").innerHTML = self.changed_appearance_description
        self.changed_damages_description += self.wood_damage_description
        self.js.document.getElementById("changed_damages_description").innerHTML = self.changed_damages_description

    def addMetal(self):
        self.changed_appearance_description += self.metal_appearance_description
        self.js.document.getElementById(
            "changed_appearance_description").innerHTML = self.changed_appearance_description
        self.changed_damages_description += self.metal_damage_description
        self.js.document.getElementById("changed_damages_description").innerHTML = self.changed_damages_description

    def hideMetalResearches(self):
        self.js.document.getElementById("choose_iron_cunduct").style.visibility = 'hidden'
        self.js.document.getElementById("choose_cuprum_cunduct").style.visibility = 'hidden'
        self.js.document.getElementById("choose_silver_cunduct").style.visibility = 'hidden'

    def showMetalResearches(self):
        self.js.document.getElementById("choose_iron_cunduct").style.visibility = 'visible'
        self.js.document.getElementById("choose_cuprum_cunduct").style.visibility = 'visible'
        self.js.document.getElementById("choose_silver_cunduct").style.visibility = 'visible'

    def fillInIronForm(self):
        self.changed_ConductTitle += self.ironConductTitle
        self.js.document.getElementById("purposes_of_conduction").innerHTML = self.changed_ConductTitle
        self.changed_ConductDescription += self.ironConductDescription
        self.js.document.getElementById("methods_conduction").innerHTML = self.changed_ConductDescription
        self.changed_restoration_program += self.restoration_program_iron
        self.js.document.getElementById("restoration_program_by_material").innerHTML = self.changed_restoration_program
        self.changed_treatments_descriptions += self.iron_treatments_descriptions
        self.js.document.getElementById(
            "treatments_descriptions_by_material").innerHTML = self.changed_treatments_descriptions
        self.changed_treatments_chemicals += self.iron_treatments_chemicals
        self.js.document.getElementById(
            "treatments_chemicals_by_material").innerHTML = self.changed_treatments_chemicals

    def fillInCuprumForm(self):
        self.changed_ConductTitle += self.cuprumConductTitle
        self.js.document.getElementById("purposes_of_conduction").innerHTML = self.changed_ConductTitle
        self.changed_ConductDescription += self.cuprumConductDescription
        self.js.document.getElementById("methods_conduction").innerHTML = self.changed_ConductDescription
        self.changed_restoration_program += self.restoration_program_cuprum
        self.js.document.getElementById("restoration_program_by_material").innerHTML = self.changed_restoration_program
        self.changed_treatments_descriptions += self.cuprum_treatments_descriptions
        self.js.document.getElementById(
            "treatments_descriptions_by_material").innerHTML = self.changed_treatments_descriptions
        self.changed_treatments_chemicals += self.cuprum_treatments_chemicals
        self.js.document.getElementById(
            "treatments_chemicals_by_material").innerHTML = self.changed_treatments_chemicals

    def fillInSilverForm(self):
        self.changed_ConductTitle += self.silverConductTitle
        self.js.document.getElementById("purposes_of_conduction").innerHTML = self.changed_ConductTitle
        self.changed_ConductDescription += self.silverConductDescription
        self.js.document.getElementById("methods_conduction").innerHTML = self.changed_ConductDescription
        self.changed_restoration_program += self.restoration_program_silver
        self.js.document.getElementById("restoration_program_by_material").innerHTML = self.changed_restoration_program
        self.changed_treatments_descriptions += self.silver_treatments_descriptions
        self.js.document.getElementById(
            "treatments_descriptions_by_material").innerHTML = self.changed_treatments_descriptions
        self.changed_treatments_chemicals += self.silver_treatments_chemicals
        self.js.document.getElementById(
            "treatments_chemicals_by_material").innerHTML = self.changed_treatments_chemicals


def connect_db():
    conn = sqlite3.connect(app.config['DATABASE'])
    conn.row_factory = sqlite3.Row
    return conn


def create_db():
    db = connect_db()
    with app.open_resource('sq_db.sql', mode='r') as f:
        db.cursor().executescript(f.read())
    db.commit()
    db.close()


def get_db():
    if not hasattr(g, 'link_db'):
        g.link_db = connect_db()
    return g.link_db


@app.teardown_appcontext
def close_db(error):
    if hasattr(g, 'link_db'):
        g.link_db.close()


@app.route('/')
def index():
    db = get_db()
    dbase = FDataBase(db)
    print(url_for('index'))
    return render_template('index.html', main_menu=dbase.getMainMenu(), web_page_title='Головна',
                           passports=dbase.get_passports_preview())


def set_name_to_image(image_name, image):
    image_name = str(image_name)
    image.filename = image_name
    displayable_image = secure_filename(image.filename)
    image.save(os.path.join(app.config['UPLOAD_FOLDER'], displayable_image))


@app.route('/add_passport', methods=['POST', 'GET'])
def add_passport():
    db = get_db()
    dbase = FDataBase(db)
    print(url_for('add_passport'))

    if request.method == 'POST':
        if len(request.form['inventory_number']) > 0 and len(request.form['acceptance_number']) > 0:

            names_and_images = {
                request.form['before_restoration_image_description']: request.files[
                    'before_restoration_image_of_object'],
                request.form['process_restoration_image_description']: request.files[
                    'process_restoration_image_of_object'],
                request.form['after_restoration_image_description']: request.files['after_restoration_image_of_object']
            }
            try:
                for image_name, image in names_and_images.copy().items():
                    set_name_to_image(image_name, image)
            except FileNotFoundError:
                flash('Назвіть фото латиницею без використання спец. символів та злитно', category='error')
                return redirect('/add_passport')

            passport_fields_input = dbase.store_passport(request.form['passport_number'],
                                                         request.form['inventory_number'],
                                                         request.form['acceptance_number'],
                                                         request.form['institution_name'],
                                                         request.form['department_name'],
                                                         request.form['definition'], request.form['typological'],
                                                         request.form['object_owner'],
                                                         request.form['author'], request.form['clarified_author'],
                                                         request.form['object_title'],
                                                         request.form['clarified_object_title'],
                                                         request.form['time_of_creation'],
                                                         request.form['clarified_time_of_creation'],
                                                         request.form['material'], request.form['clarified_material'],
                                                         request.form['technique'],
                                                         request.form['clarified_technique'],
                                                         request.form['object_size'], request.form['clarified_size'],
                                                         request.form['weight'],
                                                         request.form['clarified_weight'], request.form['reason'],
                                                         request.form['object_input_date'],
                                                         request.form['execute_restorer'],
                                                         request.form['object_output_date'],
                                                         request.form['responsible_restorer'],
                                                         request.form['origin_description'],
                                                         request.form['appearance_description'],
                                                         request.form['damages_description'],
                                                         request.form['signs_description'],
                                                         request.form['size_description'],
                                                         request.form['purposes_researches'],
                                                         request.form['methods_researches'],
                                                         request.form['executor_date_researches'],
                                                         request.form['results_researches'],
                                                         request.form['restoration_program'],
                                                         request.form['treatments_descriptions'],
                                                         request.form['treatments_chemicals'],
                                                         request.form['treatments_executor_date'],
                                                         request.form['treatments_results'],
                                                         request.form['before_restoration_image_description'],
                                                         request.files['before_restoration_image_of_object'],
                                                         request.form['process_restoration_image_description'],
                                                         request.files['process_restoration_image_of_object'],
                                                         request.form['after_restoration_image_description'],
                                                         request.files['after_restoration_image_of_object'])

            if not passport_fields_input:
                flash('Виникла, помилка публікування', category='error')
            else:
                flash('Успішно опубліковано!', category='success')
        else:
            flash("Спочатку введіть, будь ласка, інвентарний номер та дані акта приймання пам'ятки.", category='error')
    return ChangeFormsForMaterial.render(
        render_template('passport_form.html', main_menu=dbase.getMainMenu(), web_page_title='Публікація', ))


@app.route('/post/<int:id_post>')
def show_post(id_post):
    db = get_db()
    dbase = FDataBase(db)
    passport_number, inventory_number, acceptance_number, institution_name, department_name, definition, typological, object_owner, author, clarified_author, object_title, clarified_object_title, time_of_creation, clarified_time_of_creation, material, clarified_material, technique, clarified_technique, object_size, clarified_size, weight, clarified_weight, reason, object_input_date, execute_restorer, object_output_date, responsible_restorer, origin_description, appearance_description, damages_description, signs_description, size_description, purposes_researches, methods_researches, executor_date_researches, results_researches, restoration_program, treatments_descriptions, treatments_chemicals, treatments_executor_date, treatments_results, before_restoration_image_description, before_restoration_image_of_object, process_restoration_image_description, process_restoration_image_of_object, after_restoration_image_description, after_restoration_image_of_object = dbase.get_passport(
        id_post)
    build_passport(passport_number, inventory_number, acceptance_number, institution_name, department_name)
    if not inventory_number:
        abort(404)

    return render_template('post.html', main_menu=dbase.getMainMenu(), web_page_title=inventory_number,
                           passport_number=passport_number,
                           inventory_number=inventory_number, acceptance_number=acceptance_number,
                           institution_name=institution_name,
                           department_name=department_name, definition=definition, typological=typological,
                           object_owner=object_owner, author=author, clarified_author=clarified_author,
                           object_title=object_title,
                           clarified_object_title=clarified_object_title, time_of_creation=time_of_creation,
                           clarified_time_of_creation=clarified_time_of_creation, material=material,
                           clarified_material=clarified_material,
                           technique=technique, clarified_technique=clarified_technique, object_size=object_size,
                           clarified_size=clarified_size,
                           weight=weight, clarified_weight=clarified_weight, reason=reason,
                           object_input_date=object_input_date, execute_restorer=execute_restorer,
                           object_output_date=object_output_date, responsible_restorer=responsible_restorer,
                           origin_description=origin_description, appearance_description=appearance_description,
                           damages_description=damages_description, signs_description=signs_description,
                           size_description=size_description, purposes_researches=purposes_researches,
                           methods_researches=methods_researches, executor_date_researches=executor_date_researches,
                           results_researches=results_researches, restoration_program=restoration_program,
                           treatments_descriptions=treatments_descriptions, treatments_chemicals=treatments_chemicals,
                           treatments_executor_date=treatments_executor_date, treatments_results=treatments_results,
                           before_restoration_image_description=before_restoration_image_description,
                           before_restoration_image_of_object=before_restoration_image_of_object,
                           process_restoration_image_description=process_restoration_image_description,
                           process_restoration_image_of_object=process_restoration_image_of_object,
                           after_restoration_image_description=after_restoration_image_description,
                           after_restoration_image_of_object=after_restoration_image_of_object)


def build_passport(passport_number, inventory_number, acceptance_number, institution_name, department_name):
    document = Document()

    passport_identity_table = document.add_table(rows=2, cols=3)
    passport_identity_table.style = 'TableGrid'
    passport_identity_table.alignment = 2

    passport_identity_parameter = passport_identity_table.rows[0].cells
    passport_identity_parameter[0].text = '№ реставраційного паспорта:'
    passport_identity_parameter[1].text = "інвентарний № пам'ятки"
    passport_identity_parameter[2].text = 'Акт приймання'

    passport_identity_input = passport_identity_table.rows[1].cells
    passport_identity_input[0].text = f'{passport_number}'
    passport_identity_input[1].text = f'{inventory_number}'
    passport_identity_input[2].text = f'{acceptance_number}'

    ministry_title = document.add_paragraph()
    passport_type_title = document.add_paragraph()
    institution_title = document.add_paragraph()
    institution_parameter = document.add_paragraph()
    department_title = document.add_paragraph()
    department_parameter = document.add_paragraph()

    font_styles = document.styles
    font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.name = 'Times New Roman'
    font_object.size = Pt(12)

    ministry_title.add_run("\nМіністерство культури та інформаційної політики України", style='CommentsStyle').bold = True
    ministry_title.alignment = 1

    passport_type_title.add_run("ПАСПОРТ РЕСТАВРАЦІЇ ПАМ'ЯТКИ ІСТОРІЇ ТА КУЛЬТУРИ (РУХОМОЇ)\n", style='CommentsStyle').bold = True
    passport_type_title.alignment = 1

    institution_title.add_run(f'{institution_name}', style='CommentsStyle').bold = True
    institution_title.alignment = 1

    institution_parameter.add_run('назва закладу, який здійснює реставрацію')
    institution_parameter.alignment = 1

    department_title.add_run(f'{department_name}', style='CommentsStyle').bold = True
    department_title.alignment = 1

    department_parameter.add_run('назва відділу/сектору')
    department_parameter.alignment = 1

    institution_table = document.add_table(rows=2, cols=6)
    institution_table.style = 'TableGrid'
    institution_table.alignment = 1

    institution_title_parameter = institution_table.rows[0].cells
    institution_title_parameter[0].text = 'НМІУ'



    # Creating paragraph
    institution_title = document.add_paragraph()
    institution_title.alignment = 1
    # Adding content to paragraph
    institution_title.add_run('Національна академія образотворчого мистецтва і архітектури\n',
                              style='CommentsStyle')
    institution_title.underline = True
    institution_title.bold = True

    # Creating paragraph
    para = document.add_paragraph()

    # Adding content to paragraph
    underline_para = para.add_run(
        '''GeeksforGeeks is a Computer Science portal for geeks. It contains well written, well thought and well-explained computer science and programming articles, quizzes etc.''')

    # Applying undeline to true
    underline_para.underline = True

    document.add_heading(f'номер паспорта: {passport_number}', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=2)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('lisichka.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    paragraph = document.add_paragraph('Normal text, ')
    paragraph.add_run('text with emphasis.', 'Emphasis')
    paragraph.insert_paragraph_before('Lorem ipsum')

    document.add_page_break()
    document.save('filled_passport.docx')


@app.route('/download_passport')
def download_passport():
    passport = 'filled_passport.docx'
    return send_file(passport, as_attachment=True, download_name='YourFilledPassport.docx')


@app.route('/about_us')
def about_us():
    db = get_db()
    dbase = FDataBase(db)
    print(url_for('about_us'))
    return render_template('about.html', main_menu=dbase.getMainMenu(), web_page_title='Про нас')


@app.errorhandler(413)
def uploaded_image_too_large(error):
    db = get_db()
    dbase = FDataBase(db)
    return render_template('page413.html', main_menu=dbase.getMainMenu(), web_page_title='Змініть зображення'), 413


@app.errorhandler(404)
def page_not_found(error):
    db = get_db()
    dbase = FDataBase(db)
    return render_template('page404.html', main_menu=dbase.getMainMenu(), web_page_title='Не знайдено'), 404


@app.route('/contact', methods=['POST', 'GET'])
def contact():
    db = get_db()
    dbase = FDataBase(db)
    if request.method == 'POST':
        if len(request.form['user_name']) > 2:
            flash('Ваше повідомлення відправлено. Ми вдосконалюємося.', category='success')
        else:
            flash('На жаль, виникла помилка відправлення', category='error')

    return render_template('contact.html', main_menu=dbase.getMainMenu(), web_page_title='Повідомити')


if __name__ == '__main__':
    app.run(debug=True)
    app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
    app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
